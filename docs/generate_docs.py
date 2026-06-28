#!/usr/bin/env python3
"""
Generate Word document with ABCT system documentation.
Creates a comprehensive breakdown of all components.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_documentation():
    """Generate the ABCT documentation Word document."""
    doc = Document()

    # Title
    title = doc.add_heading('ABCT - Crypto Portfolio Tracker', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Technical Documentation & Component Breakdown')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. System Overview',
        '2. Technology Stack',
        '3. Backend Components',
        '   3.1 Main Application (main.py)',
        '   3.2 Configuration (config.py)',
        '   3.3 Database Layer (database.py)',
        '   3.4 API Routers',
        '   3.5 Services',
        '   3.6 Utilities',
        '4. Frontend Components',
        '5. External API Integrations',
        '6. Database Schema',
        '7. Security Considerations',
        '8. Deployment Guide'
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # Section 1: System Overview
    doc.add_heading('1. System Overview', level=1)
    doc.add_paragraph(
        'ABCT (A Better Crypto Tracker) is a self-hosted cryptocurrency portfolio tracking '
        'application that aggregates data from multiple blockchains, exchanges, and DeFi protocols. '
        'It provides a unified dashboard for monitoring crypto holdings, staking rewards, NFT collections, '
        'and historical portfolio performance.'
    )

    doc.add_heading('Key Features', level=2)
    features = [
        'Multi-blockchain wallet tracking (Cardano, Bitcoin, Ethereum)',
        'Exchange integration (Coinbase)',
        'DeFi staking position monitoring',
        'NFT collection management with floor price tracking',
        'Portfolio value history with interactive charts',
        'Daily automated snapshots',
        'Privacy mode for sensitive data',
        'Responsive dark-themed UI'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # Section 2: Technology Stack
    doc.add_heading('2. Technology Stack', level=1)

    doc.add_heading('Backend', level=2)
    backend_tech = [
        ('Python 3.9+', 'Core programming language'),
        ('FastAPI', 'Modern async web framework for REST APIs'),
        ('Uvicorn', 'ASGI server for running FastAPI'),
        ('SQLite', 'Lightweight embedded database'),
        ('aiosqlite', 'Async SQLite driver'),
        ('httpx', 'Async HTTP client for API requests'),
        ('PyJWT', 'JWT token handling for Coinbase auth'),
        ('cryptography', 'Cryptographic operations')
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technology'
    hdr_cells[1].text = 'Purpose'
    for tech, purpose in backend_tech:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = purpose

    doc.add_paragraph()
    doc.add_heading('Frontend', level=2)
    frontend_tech = [
        ('HTML5', 'Page structure'),
        ('CSS3', 'Styling with CSS variables for theming'),
        ('JavaScript (ES6+)', 'Client-side interactivity'),
        ('Chart.js 4.4.1', 'Portfolio history visualization'),
        ('Fetch API', 'REST API communication')
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technology'
    hdr_cells[1].text = 'Purpose'
    for tech, purpose in frontend_tech:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = purpose

    doc.add_page_break()

    # Section 3: Backend Components
    doc.add_heading('3. Backend Components', level=1)

    doc.add_heading('3.1 Main Application (main.py)', level=2)
    doc.add_paragraph(
        'The main entry point for the FastAPI application. Responsible for:'
    )
    main_responsibilities = [
        'Application lifecycle management (startup/shutdown)',
        'Database initialization',
        'Router registration',
        'Static file serving for frontend',
        'Background task scheduling (snapshots, NFT price collection)',
        'Health check and status endpoints'
    ]
    for item in main_responsibilities:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2 Configuration (config.py)', level=2)
    doc.add_paragraph(
        'Centralized configuration management using environment variables:'
    )
    config_vars = [
        ('BLOCKFROST_API_KEY', 'Cardano blockchain API access'),
        ('CEXPLORER_API_KEY', 'Cardano staking data'),
        ('COINBASE_API_KEY_NAME', 'Coinbase exchange integration'),
        ('COINBASE_API_PRIVATE_KEY', 'Coinbase API authentication'),
        ('ETHERSCAN_API_KEY', 'Ethereum blockchain data')
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Purpose'
    for var, purpose in config_vars:
        row_cells = table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = purpose

    doc.add_heading('3.3 Database Layer (database.py)', level=2)
    doc.add_paragraph(
        'Async SQLite database operations using aiosqlite. Manages all persistent data storage '
        'including wallets, balances, native assets, portfolio snapshots, NFT prices, and caching.'
    )

    doc.add_heading('Key Functions', level=3)
    db_functions = [
        ('init_db()', 'Initialize database schema and tables'),
        ('add_wallet()', 'Add new wallet to tracking'),
        ('get_all_wallets()', 'Retrieve all tracked wallets'),
        ('update_wallet_balance()', 'Update wallet balance data'),
        ('save_portfolio_snapshot()', 'Store daily portfolio value'),
        ('get_portfolio_history()', 'Retrieve historical snapshots'),
        ('save_nft_floor_price()', 'Store NFT collection prices'),
        ('get_cache() / set_cache()', 'Generic caching system')
    ]
    for func, desc in db_functions:
        p = doc.add_paragraph()
        p.add_run(func).bold = True
        p.add_run(f' - {desc}')

    doc.add_page_break()

    doc.add_heading('3.4 API Routers', level=2)

    routers = [
        ('wallets.py', '/wallets', [
            'GET / - List all wallets',
            'POST / - Add new wallet',
            'GET /{id} - Get wallet details',
            'DELETE /{id} - Remove wallet',
            'POST /{id}/refresh - Refresh wallet data'
        ]),
        ('portfolio.py', '/portfolio', [
            'GET /summary - Portfolio overview with totals',
            'GET /assets - All native assets across wallets',
            'GET /history - Historical portfolio values',
            'POST /snapshot - Create manual snapshot',
            'POST /history/generate - Generate historical data'
        ]),
        ('prices.py', '/prices', [
            'GET / - Current prices for tracked assets',
            'GET /history - Historical price data'
        ]),
        ('defi.py', '/defi', [
            'GET /staking - Cardano staking positions',
            'GET /staking/{stake_address} - Specific stake account'
        ]),
        ('exchanges.py', '/exchanges', [
            'GET /balances - Exchange holdings',
            'GET /coinbase/accounts - Coinbase account details'
        ]),
        ('nfts.py', '/nfts', [
            'GET / - All NFTs with values',
            'GET /summary - NFTs grouped by collection',
            'GET /collection/{policy_id} - Collection details',
            'GET /prices/status - Price collection coverage',
            'POST /prices/collect - Trigger price collection',
            'POST /refresh - Force refresh NFT data'
        ])
    ]

    for filename, prefix, endpoints in routers:
        doc.add_heading(f'{filename} ({prefix}/*)', level=3)
        for endpoint in endpoints:
            doc.add_paragraph(endpoint, style='List Bullet')

    doc.add_page_break()

    doc.add_heading('3.5 Services', level=2)

    services = [
        ('cardano.py', 'CardanoService',
         'Interfaces with Blockfrost API for Cardano blockchain data. Fetches wallet balances, '
         'native assets (tokens/NFTs), stake address information, and transaction history.'),
        ('bitcoin.py', 'BitcoinService',
         'Uses Blockstream API for Bitcoin blockchain data. Retrieves wallet balances and '
         'transaction information for BTC addresses.'),
        ('ethereum.py', 'EthereumService',
         'Connects to Etherscan API for Ethereum data. Fetches ETH balances and ERC-20 token holdings.'),
        ('pricing.py', 'PricingService',
         'Aggregates cryptocurrency prices from CoinGecko free API. Provides current prices, '
         'historical price data, and multi-currency support.'),
        ('defi.py', 'DeFiService',
         'Fetches Cardano staking information from CExplorer API. Retrieves pool delegations, '
         'rewards history, and APY calculations.'),
        ('coinbase.py', 'CoinbaseService',
         'Integrates with Coinbase Advanced Trade API using JWT authentication. Fetches '
         'account balances, portfolio value, and transaction history.'),
        ('nft.py', 'NFTService',
         'Manages NFT data collection and caching. Uses TapTools for floor prices, Koios and '
         'Blockfrost as fallbacks. Implements incremental price collection to handle rate limits.'),
        ('snapshot.py', 'SnapshotService',
         'Handles daily portfolio snapshots. Calculates total portfolio value from all sources, '
         'creates snapshots at 12:00 PM CT, and supports historical data generation.')
    ]

    for filename, classname, description in services:
        doc.add_heading(f'{filename} - {classname}', level=3)
        doc.add_paragraph(description)

    doc.add_page_break()

    doc.add_heading('3.6 Utilities', level=2)

    doc.add_heading('address.py', level=3)
    doc.add_paragraph(
        'Address validation and detection utilities. Provides functions to identify blockchain '
        'type from wallet address format (Cardano addr1/stake1, Bitcoin bc1/1/3, Ethereum 0x).'
    )

    doc.add_page_break()

    # Section 4: Frontend Components
    doc.add_heading('4. Frontend Components', level=1)

    doc.add_heading('index.html', level=2)
    doc.add_paragraph(
        'Main HTML template with responsive layout. Features collapsible sections, '
        'modal dialogs for wallet management, and Chart.js integration for portfolio history.'
    )

    doc.add_heading('Key Sections', level=3)
    sections = [
        'Header with privacy toggle and refresh controls',
        'Portfolio Summary with blockchain breakdown',
        'Portfolio Value History Chart (7d/4w/3m)',
        'Staking Positions (Cardano DeFi)',
        'Exchange Holdings (Coinbase)',
        'NFT Collection with floor prices',
        'Add Wallet Modal'
    ]
    for section in sections:
        doc.add_paragraph(section, style='List Bullet')

    doc.add_heading('app.js', level=2)
    doc.add_paragraph(
        'JavaScript application logic handling all API interactions, UI updates, and user events.'
    )

    doc.add_heading('Key Functions', level=3)
    js_functions = [
        ('loadPortfolio()', 'Fetch and display portfolio summary'),
        ('loadPortfolioHistory()', 'Fetch and render history chart'),
        ('renderPortfolioChart()', 'Create Chart.js visualization'),
        ('loadStakingPositions()', 'Display Cardano staking data'),
        ('loadExchangeBalances()', 'Show Coinbase holdings'),
        ('loadNFTs()', 'Display NFT collection'),
        ('addWallet()', 'Submit new wallet to backend'),
        ('togglePrivacyMode()', 'Show/hide sensitive values'),
        ('initCollapsibleSections()', 'Setup collapsible UI')
    ]
    for func, desc in js_functions:
        p = doc.add_paragraph()
        p.add_run(func).bold = True
        p.add_run(f' - {desc}')

    doc.add_heading('styles.css', level=2)
    doc.add_paragraph(
        'Dark-themed CSS with CSS custom properties for easy theming. Features responsive '
        'grid layouts, smooth animations, and privacy mode blur effects.'
    )

    doc.add_page_break()

    # Section 5: External API Integrations
    doc.add_heading('5. External API Integrations', level=1)

    apis = [
        ('Blockfrost', 'https://blockfrost.io', 'Cardano',
         'Primary Cardano blockchain API. Provides wallet balances, native assets, stake info.'),
        ('Blockstream', 'https://blockstream.info', 'Bitcoin',
         'Free Bitcoin API. No authentication required. Wallet balances and transactions.'),
        ('Etherscan', 'https://etherscan.io', 'Ethereum',
         'Ethereum blockchain explorer API. ETH balances and ERC-20 tokens.'),
        ('CoinGecko', 'https://coingecko.com', 'Pricing',
         'Free cryptocurrency price API. Current and historical prices.'),
        ('CExplorer', 'https://cexplorer.io', 'Cardano DeFi',
         'Cardano staking data. Pool information, rewards, APY.'),
        ('TapTools', 'https://taptools.io', 'Cardano NFTs',
         'NFT floor prices and collection metadata. Rate-limited free tier.'),
        ('Koios', 'https://koios.rest', 'Cardano',
         'Free Cardano API. Fallback for collection metadata.'),
        ('Coinbase', 'https://coinbase.com', 'Exchange',
         'Advanced Trade API with JWT auth. Account balances and portfolio.')
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'API'
    hdr_cells[1].text = 'URL'
    hdr_cells[2].text = 'Purpose'
    hdr_cells[3].text = 'Notes'
    for api, url, purpose, notes in apis:
        row_cells = table.add_row().cells
        row_cells[0].text = api
        row_cells[1].text = url
        row_cells[2].text = purpose
        row_cells[3].text = notes

    doc.add_page_break()

    # Section 6: Database Schema
    doc.add_heading('6. Database Schema', level=1)

    tables = [
        ('wallets', [
            ('id', 'INTEGER PRIMARY KEY', 'Auto-increment ID'),
            ('address', 'TEXT UNIQUE', 'Wallet address'),
            ('blockchain', 'TEXT', 'cardano/bitcoin/ethereum'),
            ('label', 'TEXT', 'User-defined label'),
            ('created_at', 'TIMESTAMP', 'Creation time'),
            ('updated_at', 'TIMESTAMP', 'Last update')
        ]),
        ('balances', [
            ('id', 'INTEGER PRIMARY KEY', 'Auto-increment ID'),
            ('wallet_id', 'INTEGER FK', 'Reference to wallets'),
            ('amount', 'TEXT', 'Balance amount'),
            ('unit', 'TEXT', 'Currency unit'),
            ('updated_at', 'TIMESTAMP', 'Last update')
        ]),
        ('native_assets', [
            ('id', 'INTEGER PRIMARY KEY', 'Auto-increment ID'),
            ('wallet_id', 'INTEGER FK', 'Reference to wallets'),
            ('asset_id', 'TEXT', 'Full asset identifier'),
            ('policy_id', 'TEXT', 'Policy ID (NFT collection)'),
            ('asset_name', 'TEXT', 'Human-readable name'),
            ('quantity', 'TEXT', 'Amount held')
        ]),
        ('portfolio_snapshots', [
            ('id', 'INTEGER PRIMARY KEY', 'Auto-increment ID'),
            ('snapshot_date', 'DATE UNIQUE', 'Snapshot date'),
            ('snapshot_time', 'TEXT', 'Time of snapshot'),
            ('total_value_usd', 'REAL', 'Total portfolio USD value'),
            ('ada_amount/price', 'REAL', 'ADA holdings and price'),
            ('btc_amount/price', 'REAL', 'BTC holdings and price'),
            ('eth_amount/price', 'REAL', 'ETH holdings and price'),
            ('*_value_usd', 'REAL', 'Component values')
        ]),
        ('nft_floor_prices', [
            ('id', 'INTEGER PRIMARY KEY', 'Auto-increment ID'),
            ('policy_id', 'TEXT', 'NFT collection policy'),
            ('collection_name', 'TEXT', 'Collection name'),
            ('floor_price_ada', 'REAL', 'Floor price in ADA'),
            ('listings', 'INTEGER', 'Active listings count'),
            ('supply', 'INTEGER', 'Total supply'),
            ('source', 'TEXT', 'Data source'),
            ('fetched_at', 'TIMESTAMP', 'Fetch time')
        ]),
        ('cache', [
            ('key', 'TEXT PRIMARY KEY', 'Cache key'),
            ('value', 'TEXT', 'JSON serialized data'),
            ('expires_at', 'TIMESTAMP', 'Expiration time')
        ])
    ]

    for table_name, columns in tables:
        doc.add_heading(table_name, level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Column'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Description'
        for col, col_type, desc in columns:
            row_cells = table.add_row().cells
            row_cells[0].text = col
            row_cells[1].text = col_type
            row_cells[2].text = desc
        doc.add_paragraph()

    doc.add_page_break()

    # Section 7: Security Considerations
    doc.add_heading('7. Security Considerations', level=1)

    doc.add_heading('API Key Management', level=2)
    security_items = [
        'All API keys stored in .env file (not committed to version control)',
        'Environment variables loaded at runtime only',
        'Coinbase uses JWT authentication with short-lived tokens',
        'No sensitive data logged or exposed in error messages'
    ]
    for item in security_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Data Privacy', level=2)
    privacy_items = [
        'Privacy mode blurs all financial values in UI',
        'Local SQLite database - data never leaves your machine',
        'No analytics or telemetry',
        'Read-only blockchain access (cannot move funds)'
    ]
    for item in privacy_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Network Security', level=2)
    network_items = [
        'Server binds to localhost (127.0.0.1) by default',
        'All external API calls use HTTPS',
        'No incoming connections required'
    ]
    for item in network_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Section 8: Deployment Guide
    doc.add_heading('8. Deployment Guide', level=1)

    doc.add_heading('Prerequisites', level=2)
    prereqs = [
        'Python 3.9 or higher',
        'pip (Python package manager)',
        'Git (optional, for cloning)',
        'API keys for desired services'
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')

    doc.add_heading('Quick Start', level=2)
    steps = [
        '1. Clone or copy the project to your machine',
        '2. Run the deployment script: ./Deployment/setup.sh',
        '3. Configure API keys in .env file',
        '4. Start the server: ./run.sh',
        '5. Open browser to http://localhost:8000'
    ]
    for step in steps:
        doc.add_paragraph(step)

    doc.add_heading('Required API Keys', level=2)
    doc.add_paragraph(
        'At minimum, you need a Blockfrost API key to track Cardano wallets. '
        'Other keys are optional depending on which features you want to use.'
    )

    required_keys = [
        ('Blockfrost', 'Required', 'https://blockfrost.io - Free tier available'),
        ('CExplorer', 'Recommended', 'https://cexplorer.io - For staking data'),
        ('TapTools', 'Optional', 'https://taptools.io - For NFT floor prices'),
        ('Coinbase', 'Optional', 'https://coinbase.com/settings/api - For exchange'),
        ('Etherscan', 'Optional', 'https://etherscan.io - For Ethereum')
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Service'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Sign Up'
    for service, status, url in required_keys:
        row_cells = table.add_row().cells
        row_cells[0].text = service
        row_cells[1].text = status
        row_cells[2].text = url

    # Save the document
    output_path = os.path.join(os.path.dirname(__file__), 'ABCT_Technical_Documentation.docx')
    doc.save(output_path)
    print(f'Documentation saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_documentation()
